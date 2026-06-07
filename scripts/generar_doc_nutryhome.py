#!/usr/bin/env python3
"""Genera un documento Word detallado con las caracteristicas de NutryHome."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# Paleta de colores
AZUL = RGBColor(0x1F, 0x4E, 0x79)
AZUL_CLARO = RGBColor(0x2E, 0x74, 0xB5)
GRIS = RGBColor(0x59, 0x59, 0x59)
VERDE = RGBColor(0x2E, 0x7D, 0x32)
NARANJA = RGBColor(0xC0, 0x55, 0x00)

doc = Document()

# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def set_heading_color(par, color):
    for run in par.runs:
        run.font.color.rgb = color


def add_heading(text, level=1, color=AZUL):
    h = doc.add_heading(text, level=level)
    set_heading_color(h, color)
    return h


def add_para(text, bold=False, italic=False, color=None, size=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_kv_table(rows, headers=None, widths=None):
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if headers:
        hdr = table.add_row().cells
        for i, h in enumerate(headers):
            hdr[i].paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


# ============ PORTADA ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("NutryHome")
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = AZUL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Sistema de Call Center con IA de Voz para Nutricion y Salud")
r.font.size = Pt(15)
r.font.color.rgb = AZUL_CLARO
r.italic = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run("Descripcion Tecnica y Funcional Detallada")
r2.font.size = Pt(13)
r2.font.color.rgb = GRIS

fecha = doc.add_paragraph()
fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fecha.add_run("Documento generado el " + datetime.now().strftime("%d/%m/%Y"))
rf.font.size = Pt(10)
rf.font.color.rgb = GRIS

doc.add_paragraph()

# ============ 1. RESUMEN EJECUTIVO ============
add_heading("1. Resumen Ejecutivo", level=1)
add_para(
    "NutryHome es una plataforma fullstack de call center automatizado orientada al sector "
    "de nutricion y salud. Su nucleo es una agente de voz con inteligencia artificial llamada "
    "\"Isabela\", construida sobre la plataforma ElevenLabs Conversational AI, que realiza "
    "llamadas salientes a pacientes para verificar stock de productos, confirmar pedidos, "
    "registrar reclamos y derivar casos."
)
add_para(
    "El sistema permite cargar contactos masivamente desde archivos Excel, ejecutar campanas "
    "outbound por lotes (batches), monitorear metricas operativas en tiempo real, escuchar "
    "grabaciones, leer transcripciones y resumenes traducidos al espanol, y generar reportes "
    "de productos detectados en las conversaciones. La linea de desarrollo mas reciente "
    "incorpora una integracion con WhatsApp a traves de Builderbot, actualmente en fase de "
    "implementacion parcial."
)

add_heading("Propuesta de valor", level=2, color=AZUL_CLARO)
add_bullet("Automatiza el contacto saliente con pacientes mediante voz IA, reduciendo carga del equipo humano.")
add_bullet("Centraliza metricas, grabaciones, transcripciones y reportes en un unico dashboard.")
add_bullet("Procesa contactos masivos desde Excel y los convierte en campanas ejecutables.")
add_bullet("Extrae automaticamente productos y cantidades mencionados en cada conversacion.")
add_bullet("Preparado para extender la operacion a canales de mensajeria (WhatsApp/Builderbot).")

# ============ 2. ARQUITECTURA ============
add_heading("2. Arquitectura General", level=1)
add_para(
    "El sistema se compone de dos grandes piezas: un frontend Next.js 14 (App Router) que "
    "tambien expone route handlers que actuan como proxy/orquestador, y un backend Express "
    "con base de datos PostgreSQL gestionada con Prisma. ElevenLabs provee la capa de voz y "
    "conversacion."
)

add_heading("Componentes principales", level=2, color=AZUL_CLARO)
add_kv_table(
    rows=[
        ["Frontend", "Next.js 14 + React 18 + TypeScript", "Vercel"],
        ["API/Proxy", "Route Handlers de Next.js (src/app/api)", "Vercel"],
        ["Backend", "Node.js + Express + Prisma", "Railway"],
        ["Base de datos", "PostgreSQL", "Railway"],
        ["Voz / IA conversacional", "ElevenLabs Conversational AI", "Cloud ElevenLabs"],
        ["Mensajeria (nuevo)", "WhatsApp Cloud API / Builderbot", "Externo"],
    ],
    headers=["Capa", "Tecnologia", "Despliegue"],
)

add_para(
    "Nota tecnica: el frontend real vive en la raiz del repositorio, aunque la documentacion "
    "antigua (README.md y SETUP.md) todavia hace referencia a una carpeta frontend/. El "
    "backend de produccion responde en https://nutryhome-production.up.railway.app y se usa "
    "como fallback desde los route handlers de Next.",
    italic=True, color=GRIS,
)

# ============ 3. STACK ============
add_heading("3. Stack Tecnologico", level=1)

add_heading("3.1 Frontend", level=2, color=AZUL_CLARO)
add_bullet("framework principal de la aplicacion.", bold_prefix="Next.js 14 + React 18 + TypeScript: ")
add_bullet("estilado y diseno responsive.", bold_prefix="TailwindCSS + Radix UI + lucide-react: ")
add_bullet("graficos y visualizaciones (KPIs, evolucion, donut).", bold_prefix="Recharts: ")
add_bullet("manejo de estado y autenticacion local del cliente.", bold_prefix="Zustand: ")
add_bullet("lectura y exportacion de archivos Excel.", bold_prefix="XLSX (SheetJS): ")
add_bullet("cliente HTTP hacia el backend y servicios externos.", bold_prefix="Axios: ")
add_bullet("manejo de fechas, notificaciones y animaciones.", bold_prefix="date-fns, react-hot-toast, framer-motion: ")
add_bullet("traduccion automatica de resumenes de ingles a espanol.", bold_prefix="google-translate-api-x: ")

add_heading("3.2 Backend", level=2, color=AZUL_CLARO)
add_bullet("servidor de API REST.", bold_prefix="Node.js + Express: ")
add_bullet("ORM y acceso a PostgreSQL.", bold_prefix="Prisma: ")
add_bullet("autenticacion y hashing de contrasenas.", bold_prefix="JWT + bcryptjs: ")
add_bullet("validacion de payloads de entrada.", bold_prefix="Joi + express-validator: ")
add_bullet("seguridad, CORS, rate limiting y logging.", bold_prefix="helmet, cors, express-rate-limit, morgan: ")
add_bullet("carga y parseo de archivos Excel/CSV.", bold_prefix="multer, csv-parser, xlsx: ")
add_bullet("testing.", bold_prefix="Jest + supertest: ")

# ============ 4. FUNCIONALIDADES ============
add_heading("4. Funcionalidades Principales", level=1)

funcs = [
    ("Dashboard operativo",
     "Panel central con indicadores clave: total de llamadas, duracion total y promedio, "
     "cantidad de reclamos, derivaciones, eficiencia operativa, grafico de evolucion de los "
     "ultimos 7 dias, estado operativo en vivo, donut de pacientes y tabla de las ultimas 20 "
     "llamadas con controles de grabacion."),
    ("Gestion de llamadas (Isabela)",
     "Vista de historial de llamadas con busqueda, filtros por estado, paginacion, acceso a la "
     "transcripcion completa, resumen traducido, resultados de evaluacion, notas y reproductor "
     "de audio (incluye un mini-reproductor flotante moderno)."),
    ("Campanas outbound de voz",
     "Creacion y gestion de lotes (batches) de llamadas. Permite ejecutar las campanas via "
     "ElevenLabs Batch Calling, pausar, cancelar, eliminar, sincronizar resultados desde "
     "ElevenLabs y consultar el detalle de cada lote."),
    ("Carga masiva por Excel",
     "Subida de archivos Excel con contactos. Cada fila incluye telefono, nombre del contacto y "
     "del paciente, domicilio, localidad, delegacion, hasta 5 productos con sus cantidades, "
     "observaciones y prioridad."),
    ("Reporte de productos",
     "Genera un reporte tabular de productos y cantidades detectados, extrayendo informacion "
     "tanto de las variables dinamicas como de las transcripciones de ElevenLabs. Exportable a "
     "formato Excel (XLSX)."),
    ("Mensajes WhatsApp / Builderbot",
     "Modulo nuevo con bandeja de entrada, carga de contactos, gestion de lotes y detalle de "
     "conversacion. Actualmente en gran parte en modo demo, a la espera de la conexion real con "
     "el proveedor de mensajeria."),
    ("Callbacks / Webhooks",
     "Pantalla operativa para copiar y consultar las URLs de los webhooks de ElevenLabs, "
     "WhatsApp, Builderbot y el batch dispatch, facilitando la configuracion de integraciones."),
    ("Pacientes",
     "Modulo reservado, actualmente en desarrollo."),
]
for nombre, desc in funcs:
    add_heading(nombre, level=2, color=AZUL_CLARO)
    add_para(desc)

# ============ 5. PAGINAS ============
add_heading("5. Mapa de Paginas (Frontend)", level=1)
add_kv_table(
    rows=[
        ["/", "Dashboard principal con KPIs y ultimas llamadas"],
        ["/login", "Inicio de sesion (auth local visible: admin/admin)"],
        ["/calls", "Gestion e historial de llamadas de Isabela"],
        ["/calls/campanas", "Campanas y lotes outbound"],
        ["/upload", "Carga de llamadas/contactos por Excel"],
        ["/estadisticas-isabela", "Vista de gestion de llamadas (alias)"],
        ["/reportes/productos", "Reporte de productos detectados"],
        ["/callbacks", "Documentacion y URLs de webhooks"],
        ["/pacientes", "Placeholder (en desarrollo)"],
        ["/mensajes/bandeja", "Bandeja de hilos WhatsApp/Builderbot (demo)"],
        ["/mensajes/carga", "Carga Excel para campanas WhatsApp (demo)"],
        ["/mensajes/batches", "Lotes de mensajes (demo)"],
        ["/mensajes/conversacion/[id]", "Detalle de conversacion (demo)"],
    ],
    headers=["Ruta", "Descripcion"],
)

# ============ 6. ENDPOINTS API NEXT ============
add_heading("6. Endpoints API (Next.js Route Handlers)", level=1)
add_kv_table(
    rows=[
        ["GET /api/estadisticas-isabela", "Consulta ElevenLabs, pagina conversaciones, traduce resumenes y arma el dashboard"],
        ["GET /api/reports/elevenlabs", "Genera reporte JSON/XLSX con extraccion de productos desde transcripciones"],
        ["GET /api/reports/productos", "Proxy al backend para el reporte de productos"],
        ["GET /api/conversations", "Proxy a Railway /api/elevenlabs/conversations"],
        ["GET /api/get-transcript", "Proxy para obtener transcripcion por ID"],
        ["GET /api/get-audio", "Proxy para obtener audio por ID"],
        ["POST /api/recover-historical", "Recupera historico de conversaciones de ElevenLabs"],
        ["GET /api/test", "Health check simple"],
        ["GET/POST /api/webhooks/whatsapp", "Verificacion Meta y recepcion de eventos WhatsApp"],
        ["GET/POST /api/webhooks/builderbot", "Webhook generico de Builderbot con secreto opcional"],
        ["POST /api/builderbot/batch-dispatch", "Recibe lotes WhatsApp desde Railway"],
    ],
    headers=["Endpoint", "Funcion"],
)

# ============ 7. BACKEND EXPRESS ============
add_heading("7. Endpoints del Backend (Express)", level=1)
add_para("El servidor Express monta los siguientes grupos de rutas:")
add_kv_table(
    rows=[
        ["/api/calls", "CRUD de llamadas"],
        ["/api/stats", "Overview, derivaciones, reclamos y performance"],
        ["/api/webhooks", "Webhook firmado de ElevenLabs y health/test"],
        ["/api/auth", "Login, register, me, refresh y logout (JWT)"],
        ["/api/campaigns", "Campanas, batches, carga Excel/CSV, ejecucion, sync y cancelacion"],
        ["/api/variables", "Plantillas de variables dinamicas"],
        ["/api/isabela", "Sync y listado de conversaciones, resumenes traducidos y audio"],
        ["/api/elevenlabs", "Conversaciones, detalle, audio, post-call webhook e historico"],
        ["/api/reports", "Reporte Excel de productos desde ElevenLabs"],
        ["/api/contacts/by-phone", "Hidratacion de variables de contacto por telefono (Builderbot)"],
    ],
    headers=["Grupo de rutas", "Funcion"],
)

# ============ 8. INTEGRACIONES ============
add_heading("8. Integraciones Externas", level=1)
add_bullet("nucleo de voz e IA conversacional. Provee API key, agent id, phone number id, batch calling, conversaciones, audio, transcripciones y webhooks post-call.", bold_prefix="ElevenLabs Conversational AI: ")
add_bullet("aloja el backend Express y la base de datos PostgreSQL en produccion.", bold_prefix="Railway: ")
add_bullet("aloja el frontend Next.js y sus route handlers.", bold_prefix="Vercel: ")
add_bullet("webhook preparado (GET/POST) para verificacion y recepcion de eventos.", bold_prefix="WhatsApp Cloud API / Meta: ")
add_bullet("endpoint para hidratar variables de contacto y batch dispatch de lotes WhatsApp.", bold_prefix="Builderbot: ")
add_bullet("traduccion de resumenes de conversaciones de ingles a espanol.", bold_prefix="Google Translate: ")
add_bullet("mencionado como posible destino del dispatch; sin integracion concreta implementada todavia.", bold_prefix="n8n: ")

# ============ 9. MODELO DE DATOS ============
add_heading("9. Modelo de Datos (Prisma / PostgreSQL)", level=1)
add_para("El esquema define los siguientes modelos principales:")
add_kv_table(
    rows=[
        ["Campaign", "Campana con estado, fechas, tipo y usuario creador"],
        ["Batch", "Lote de llamadas asociado a una campana, contadores y elevenLabsBatchId"],
        ["Contact", "Contacto/paciente cargado desde Excel (hasta 5 productos/cantidades, direccion, prioridad)"],
        ["OutboundCall", "Llamada saliente: estado, intentos, resultado, IDs ElevenLabs, transcript, resumen, audio"],
        ["Call", "Llamada generica (inbound/outbound) con transcript, data collection y criteria results"],
        ["Derivation", "Derivaciones asociadas a una llamada"],
        ["Complaint", "Reclamos asociados a una llamada (tipo, severidad, resuelto)"],
        ["User", "Usuarios del backend con rol y autenticacion JWT"],
        ["IsabelaConversation", "Conversacion de ElevenLabs persistida con conversationId y summary"],
    ],
    headers=["Modelo", "Descripcion"],
)

add_heading("Enumeraciones (enums) clave", level=2, color=AZUL_CLARO)
add_kv_table(
    rows=[
        ["CampaignStatus", "DRAFT, ACTIVE, PAUSED, COMPLETED, CANCELLED"],
        ["BatchStatus", "PENDING, PROCESSING, COMPLETED, FAILED, CANCELLED"],
        ["OutboundCallStatus", "PENDING, SCHEDULED, IN_PROGRESS, COMPLETED, FAILED, CANCELLED"],
        ["OutboundCallResult", "ANSWERED, NO_ANSWER, BUSY, INVALID_NUMBER, VOICEMAIL, HANGUP, ERROR"],
        ["LlamadaStatus", "PENDIENTE, EN_PROGRESO, COMPLETADA, FALLIDA, CANCELADA"],
        ["PedidoStatus", "PENDIENTE, ENVIADO, ENTREGADO, CANCELADO"],
        ["Priority", "BAJA, MEDIA, ALTA, CRITICA"],
        ["ComplaintType", "CALIDAD, SERVICIO, TECNICO, FACTURACION, OTRO"],
        ["UserRole", "ADMIN, SUPERVISOR, AGENTE, ANALISTA"],
    ],
    headers=["Enum", "Valores"],
)

# ============ 10. FLUJO OPERATIVO ============
add_heading("10. Flujo Operativo Tipico", level=1)
pasos = [
    "El operador carga un archivo Excel con los contactos/pacientes a contactar (/upload).",
    "El sistema crea una campana y uno o varios lotes (batches) con los contactos cargados.",
    "Se ejecuta el lote a traves de ElevenLabs Batch Calling, que realiza las llamadas con la agente Isabela.",
    "Durante la llamada, Isabela verifica stock, confirma productos/cantidades y registra observaciones o reclamos.",
    "Al finalizar, ElevenLabs envia el resultado (transcript, audio, resumen) via webhook post-call.",
    "El sistema persiste la conversacion, traduce el resumen al espanol y actualiza el estado de la llamada.",
    "El operador monitorea las metricas en el dashboard y consulta transcripciones, audios y reportes de productos.",
]
for i, paso in enumerate(pasos, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(paso)

# ============ 11. CAMBIOS RECIENTES ============
add_heading("11. Cambios Recientes Destacados", level=1)
add_para("El desarrollo reciente se centra en WhatsApp/Builderbot y en mejoras de la experiencia de llamadas:")
add_bullet("GET /api/contacts/by-phone para hidratar variables en Builderbot.")
add_bullet("API batch-dispatch + webhooks WhatsApp/Builderbot + UI de Callbacks.")
add_bullet("Batch de WhatsApp con variables tipo ElevenLabs y despacho opcional.")
add_bullet("Nueva pagina Callbacks (ElevenLabs + WhatsApp/Builderbot).")
add_bullet("Modulo de Mensajes: layout, submenu, carga, lotes, bandeja y detalle.")
add_bullet("Mini reproductor flotante moderno en la vista de llamadas.")
add_bullet("Dashboard con ultimas 20 llamadas y controles de grabacion.")

# ============ 12. ESTADO Y OBSERVACIONES ============
add_heading("12. Estado Actual y Observaciones Tecnicas", level=1)

add_heading("Componentes maduros (produccion)", level=2, color=VERDE)
add_bullet("Llamadas ElevenLabs/Isabela: transcripciones, audio y resumenes.")
add_bullet("Campanas outbound: carga Excel, batches y ejecucion.")
add_bullet("Reportes de productos y dashboard de metricas.")

add_heading("En desarrollo / pendiente", level=2, color=NARANJA)
add_bullet("Integracion WhatsApp/Builderbot: endpoints y UI listos, pero el envio y la persistencia real aun no estan completos.")
add_bullet("Modulo de Pacientes: solo placeholder.")
add_bullet("Validacion HMAC del webhook de WhatsApp en Next requiere acceso al raw body (pendiente).")

add_heading("Puntos a revisar / deuda tecnica", level=2, color=GRIS)
add_bullet("La autenticacion visible en el frontend es local (admin/admin); el backend si tiene JWT real.")
add_bullet("Existen desalineaciones en codigo antiguo del backend (referencias a variableTemplate, campos email en OutboundCall, estados COMPLETADA/FALLIDA) que no coinciden del todo con el schema.prisma actual.")
add_bullet("README.md y SETUP.md todavia mencionan una carpeta frontend/ inexistente; el frontend vive en la raiz.")

# Cierre
doc.add_paragraph()
cierre = doc.add_paragraph()
cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = cierre.add_run("--- Fin del documento ---")
rc.italic = True
rc.font.color.rgb = GRIS

out = "/Users/ralborta/NutryHome/NutryHome_Descripcion_Detallada.docx"
doc.save(out)
print("Documento generado:", out)
