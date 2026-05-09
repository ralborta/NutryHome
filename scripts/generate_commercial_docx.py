#!/usr/bin/env python3
"""Genera documento Word comercial para NutryHome (uso interno)."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/ralborta/NutryHome/docs/NutryHome_Descripcion_Comercial.docx"


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    # Español: idioma del documento
    sect = doc.sections[0]._sectPr
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "es-ES")
    sect.append(lang)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold_first=False):
    p = doc.add_paragraph()
    if bold_first and ":" in text:
        parts = text.split(":", 1)
        run = p.add_run(parts[0] + ":")
        run.bold = True
        p.add_run(parts[1])
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def main():
    doc = Document()
    set_doc_defaults(doc)

    # Portada simple
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("NutryHome")
    r.bold = True
    r.font.size = Pt(28)
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(
        "Plataforma integral de gestión y analítica de llamadas\n"
        "para operaciones de contact center y seguimiento nutricional"
    )
    sr.font.size = Pt(14)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Documento de posicionamiento comercial\n"
        "Versión orientada a presentación ante clientes y partners"
    ).italic = True

    doc.add_page_break()

    # --- Resumen ejecutivo ---
    add_heading(doc, "1. Resumen ejecutivo", 1)
    doc.add_paragraph(
        "NutryHome es una plataforma de software empresarial diseñada para organizaciones "
        "que necesitan centralizar, medir y optimizar sus interacciones telefónicas, "
        "especialmente en contextos de call center, campañas de salud o nutrición, "
        "y seguimiento de pacientes o clientes. La solución combina un panel de control "
        "en tiempo casi real, gestión operativa de llamadas entrantes y salientes, "
        "integración nativa con inteligencia conversacional (ElevenLabs) y reporting "
        "avanzado exportable a Excel, reduciendo fricción operativa y elevando la "
        "visibilidad del negocio."
    )

    add_heading(doc, "2. Propuesta de valor", 1)
    bullets = [
        "Un solo lugar para ver el rendimiento del contact center: volumen, duración, éxito, derivaciones y señales de reclamo.",
        "Orquestación de campañas y lotes (batches) de llamadas salientes con contactos enriquecidos (datos de paciente, domicilio, productos y prioridades).",
        "Conexión directa con ElevenLabs: recepción de webhooks, almacenamiento de transcripciones, resúmenes, variables dinámicas y audio cuando esté disponible.",
        "Experiencia web moderna, responsive y preparada para equipos distribuidos (supervisión, agentes y analistas).",
        "Seguridad de acceso mediante autenticación JWT, roles diferenciados y buenas prácticas en API (validación, límites de peticiones, CORS).",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    add_heading(doc, "3. Público objetivo y casos de uso", 1)
    doc.add_paragraph(
        "La plataforma encaja de forma natural en laboratorios, distribuidoras de nutrición "
        "clínica, operadores de telemedicina, BPOs con vertical de salud y cualquier empresa "
        "que combine llamadas masivas salientes con necesidad de trazabilidad por llamada."
    )
    use_cases = [
        "Campañas de verificación de stock o seguimiento de pedidos con listas cargadas desde Excel.",
        "Supervisión de agentes de voz conversacional (Isabela / NutryHome) con lectura de resúmenes y criterios de evaluación.",
        "Detección asistida de derivaciones a especialistas y de posibles reclamos a partir del contenido de la conversación.",
        "Reporting de productos mencionados o acordados en llamada, con exportación a hoja de cálculo para finanzas o logística.",
    ]
    for u in use_cases:
        doc.add_paragraph(u, style="List Bullet")

    add_heading(doc, "4. Módulos y funcionalidades detalladas", 1)

    add_heading(doc, "4.1 Dashboard (vista general)", 2)
    doc.add_paragraph(
        "La pantalla principal ofrece una radiografía del call center. Integra métricas "
        "calculadas a partir de las conversaciones NutryHome/Isabela: total de llamadas, "
        "duración promedio, tasa de éxito, conteo orientativo de derivaciones y de "
        "posibles reclamos (analizando resúmenes y datos de evaluación). Incluye "
        "visualizaciones con gráficos (llamadas, derivaciones), lista de llamadas recientes "
        "y bloques de rendimiento, con opciones de actualización y enfoque en datos "
        "accionables para mandos medios y dirección."
    )

    add_heading(doc, "4.2 Gestión de llamadas", 2)
    doc.add_paragraph(
        "Módulo operativo para administrar el ciclo de vida de las llamadas. Permite "
        "trabajar en modo saliente (outbound) u orientado a entrantes según configuración "
        "de la operación, con gestión de lotes (batches): creación, seguimiento de estado, "
        "detalle de contactos asociados, programación, pruebas de llamada y acciones de "
        "mantenimiento (incluida eliminación controlada de lotes con confirmación). "
        "Incluye generación de reportes de productos vinculados a transcripciones y "
        "descarga en formato Excel para análisis externo."
    )

    add_heading(doc, "4.3 Campañas, lotes y contactos", 2)
    doc.add_paragraph(
        "El modelo de negocio de NutryHome se articula en campañas y batches. Cada campaña "
        "puede tener nombre, descripción, fechas, estado (borrador, activa, pausada, "
        "completada, cancelada) y tipo. Los lotes agrupan contactos y llevan contadores "
        "de llamadas totales, completadas y fallidas, con identificador opcional de batch "
        "en ElevenLabs para sincronización."
    )
    doc.add_paragraph(
        "Cada contacto puede almacenar: datos del cuidador y del paciente, teléfono, "
        "domicilio, localidad, delegación, fecha de envío, hasta cinco productos con sus "
        "cantidades, observaciones, prioridad, estado del pedido y estado/evolución de la "
        "llamada (pendiente, en progreso, completada, fallida, cancelada). Este nivel de "
        "detalle permite personalizar scripts y priorizar callbacks."
    )

    add_heading(doc, "4.4 Llamadas salientes (Outbound) y registro central (Call)", 2)
    doc.add_paragraph(
        "Las llamadas salientes guardan teléfono, nombre, estado, reintentos, fechas "
        "programadas y ejecutadas, resultado tipificado (contestada, no contesta, ocupado, "
        "número inválido, buzón, colgado, error, etc.), notas, duración y vínculo con "
        "ElevenLabs (ID de llamada, variables JSON, reintentos). Se almacenan resumen "
        "de conversación, transcripción completa, variables dinámicas y URL de audio."
    )
    doc.add_paragraph(
        "El registro unificado de llamada (Call) conserva identificador único, fecha, "
        "teléfono, duración, transcripción, recolección de datos estructurados (JSON), "
        "resultados de criterios, estado y tipo (entrante/saliente). Sobre esta base se "
        "relacionan derivaciones y reclamos para analítica posterior."
    )

    add_heading(doc, "4.5 Derivaciones y reclamos", 2)
    doc.add_paragraph(
        "Las derivaciones capturan motivo, descripción opcional y prioridad, vinculadas a "
        "cada llamada. Los reclamos incluyen tipo (calidad, servicio, técnico, facturación, "
        "otro), descripción, severidad y seguimiento de resolución. Esta estructura permite "
        "informes de estadísticas por motivo y severidad, alineados con KPIs de calidad y "
        "experiencia del cliente."
    )

    add_heading(doc, "4.6 Integración ElevenLabs (webhooks y conversaciones)", 2)
    doc.add_paragraph(
        "El backend expone endpoints de webhook para ingestar eventos de ElevenLabs de "
        "forma verificada (secreto configurable). Flujos adicionales permiten listar y "
        "consultar conversaciones, recuperar audio, sincronizar datos históricos y "
        "actualizar lotes desde la información almacenada en ElevenLabs. Las conversaciones "
        "de Isabela pueden persistirse con identificador de ElevenLabs y resumen traducido, "
        "reforzando el historial consultable desde la aplicación."
    )

    add_heading(doc, "4.7 Mensajes y conversaciones (interfaz de usuario)", 2)
    doc.add_paragraph(
        "La sección de Mensajes ofrece un historial tipo bandeja con pestañas para "
        "diferenciar flujos (incluida la experiencia Isabela), búsqueda, filtros por fecha "
        "y estado, y visualización enriquecida con transcripción, reproducción de audio "
        "cuando aplique, etiquetas y metadatos de satisfacción o evaluación. Está pensada "
        "para equipos de calidad y retención que necesitan auditar conversaciones sin "
        "cambiar de herramienta."
    )

    add_heading(doc, "4.8 Carga masiva de llamadas (Excel)", 2)
    doc.add_paragraph(
        "Flujo dedicado de carga de archivos Excel (.xlsx) mediante arrastrar y soltar o "
        "selector de archivo, con nombre de lote obligatorio. El sistema envía el fichero al "
        "API de campañas para crear o alimentar lotes y contactos de forma masiva, "
        "mostrando resultado (éxitos, totales y errores de validación). Reduce carga "
        "manual y acelera el lanzamiento de campañas."
    )

    add_heading(doc, "4.9 Estadísticas y reportes NutryHome / Isabela", 2)
    doc.add_paragraph(
        "Vista especializada (Llamadas NutryHome) con listado de conversaciones, KPIs "
        "agregados (llamadas totales, minutos, etc.), acceso a transcripciones en modal, "
        "reproducción de audio y enlaces a reportes de productos. Los reportes pueden "
        "presentarse en JSON en pantalla y exportarse a Excel con columnas de contacto, "
        "paciente, ubicación, fecha de llamada y desglose por hasta cinco líneas de "
        "producto/cantidad, incluyendo campos normalizados de unidades y presentación "
        "cuando el dato está disponible."
    )

    add_heading(doc, "4.10 Plantillas y variables dinámicas", 2)
    doc.add_paragraph(
        "La API contempla gestión de plantillas de variables (listado, detalle, alta, edición, "
        "borrado) y consulta de variables disponibles para alinear mensajes y campañas con "
        "los datos reales del CRM operativo documentados en el proyecto."
    )

    add_heading(doc, "4.11 Autenticación y roles", 2)
    doc.add_paragraph(
        "Inicio de sesión por email y contraseña con emisión de JWT, registro de usuarios, "
        "consulta del perfil actual, refresh y cierre de sesión. Los roles contemplados "
        "(administrador, supervisor, agente, analista) permiten adaptar el despliegue a "
        "jerarquías reales de contact center y cumplimiento de mínimo privilegio."
    )

    add_heading(doc, "4.12 API REST y estadísticas avanzadas", 2)
    doc.add_paragraph(
        "Además de autenticación y llamadas CRUD, la API ofrece: resumen global (overview), "
        "top de derivaciones, estadísticas de reclamos, métricas de desempeño (performance), "
        "reportes de productos, rutas de campañas y batches extensas (ejecución, estado, "
        "sincronización con ElevenLabs, cancelación, resúmenes), salud de webhooks y "
        "herramientas de diagnóstico controladas. Esto habilita integraciones futuras, "
        "ETL hacia BI corporativo o automatizaciones."
    )

    add_heading(doc, "5. Experiencia de usuario y diseño", 1)
    doc.add_paragraph(
        "Interfaz construida con Next.js y Tailwind CSS, animaciones fluidas, iconografía "
        "clara y navegación lateral con accesos directos a Dashboard, Gestión de llamadas, "
        "Mensajes, Carga de llamadas y Estadísticas NutryHome. El enfoque responsive y las "
        "referencias del proyecto a accesibilidad (WCAG 2.1) y preparación PWA refuerzan "
        "el mensaje de producto maduro para entornos corporativos mixtos (oficina y campo)."
    )

    add_heading(doc, "6. Infraestructura y despliegue", 1)
    doc.add_paragraph(
        "Arquitectura típica en producción: frontend alojado en Vercel para CDN global y "
        "despliegue continuo; backend y base de datos PostgreSQL en Railway; base relacional "
        "gestionada con Prisma para migraciones y consistencia. Este stack facilita "
        "escalado, revisiones de esquema y operación con costes predecibles."
    )

    add_heading(doc, "7. Seguridad y gobernanza", 1)
    doc.add_paragraph(
        "Validación de entradas (Joi), límites de frecuencia de peticiones, CORS configurado, "
        "contraseñas tratadas de forma segura en el backend y tokens JWT para sesiones. "
        "Los webhooks de terceros se validan mediante secreto compartido. La combinación "
        "responde a expectativas habituales de TI en proyectos B2B."
    )

    add_heading(doc, "8. Beneficios mensurables para el negocio", 1)
    benefits = [
        "Reducción del tiempo de supervisión gracias a dashboards y listados filtrables.",
        "Mayor trazabilidad por llamada, contacto y campaña, con impacto en auditorías y calidad.",
        "Menor dependencia de hojas sueltas: exportaciones estructuradas a Excel.",
        "Base lista para escalar volumen de llamadas y nuevas integraciones vía API.",
    ]
    for b in benefits:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "9. Cierre", 1)
    doc.add_paragraph(
        "NutryHome no es solo un CRM de llamadas: es una capa operativa y analítica que "
        "conecta la voz del cliente (humana o asistida por IA), los datos de campaña y los "
        "indicadores que la dirección necesita para decidir. Este documento resume las "
        "capacidades implementadas en la plataforma tal como están reflejadas en el "
        "producto actual, con un tono preparado para presentaciones comerciales, licitaciones "
        "o partnerships tecnológicos."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("— Fin del documento —").italic = True

    doc.save(OUT)
    print(f"Generado: {OUT}")


if __name__ == "__main__":
    main()
